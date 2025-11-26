import os, re, json, pymysql, torch, whisper, dateparser
from datetime import datetime, timedelta
from typing import List, Optional
from pydantic import BaseModel
from langchain_openai import ChatOpenAI
from langchain.prompts import PromptTemplate
from prompts.meeting_summary_prompt import meeting_summary_prompt

# ===================== 설정 =====================
from dotenv import load_dotenv
import os

load_dotenv()
api_key = os.getenv("OPENAI_API_KEY")

DB_HOST, DB_USER, DB_PASSWORD, DB_NAME, DB_PORT = (
    "112.175.29.231", "cheolwoo", "1234", "meeting_summary2", 33067
)

# ===================== Pydantic 구조 =====================
class ActionItem(BaseModel):
    name: str
    task: str
    due: Optional[str]

class MeetingSummary(BaseModel):
    topic_summary: str
    content_summary: str
    decisions: List[str]
    action_items: List[ActionItem]


# ===================== 날짜 정규화 함수 (강화 + 연도 보정) =====================
def normalize_due(due_text: Optional[str], base_dt: datetime) -> Optional[str]:
    """LLM이 반환한 due 문자열을 문맥기반 실제 날짜로 변환"""
    if not due_text:
        return None
    s = str(due_text).strip()
    if s in {"미정", "불명", "null", "없음", ""}:
        return None
    if s in {"오늘", "오늘 중"}:
        return base_dt.strftime("%Y-%m-%d")
    if s.startswith("내일"):
        return (base_dt + timedelta(days=1)).strftime("%Y-%m-%d")
    if s.startswith("모레"):
        return (base_dt + timedelta(days=2)).strftime("%Y-%m-%d")

    # 직접 YYYY-MM-DD 형식인 경우
    if re.match(r"\d{4}-\d{2}-\d{2}", s):
        # 🔹 과거 연도면 자동으로 올해로 교체
        y, m, d = map(int, s.split("-"))
        if y < base_dt.year:
            s = f"{base_dt.year}-{m:02d}-{d:02d}"
        return s

    # 이번주 / 다음주 패턴
    week_map = {
        "이번주 월요일": 0, "이번주 화요일": 1, "이번주 수요일": 2,
        "이번주 목요일": 3, "이번주 금요일": 4,
        "다음주 월요일": 7, "다음주 화요일": 8, "다음주 수요일": 9,
        "다음주 목요일": 10, "다음주 금요일": 11,
    }
    for k, d in week_map.items():
        if k in s:
            return (base_dt + timedelta(days=d)).strftime("%Y-%m-%d")

    # 일반 자연어 날짜 파싱 (연도 보정)
    parsed = dateparser.parse(
        s,
        languages=["ko"],
        settings={
            "RELATIVE_BASE": base_dt,
            "PREFER_DATES_FROM": "future"  # ✅ 미래 날짜 우선
        },
    )
    if parsed:
        # 🔹 연도 보정: 과거 연도면 올해로 덮어쓰기
        if parsed.year < base_dt.year:
            parsed = parsed.replace(year=base_dt.year)
        return parsed.strftime("%Y-%m-%d")

    return None


# ===================== 안전한 JSON 파싱 =====================
def safe_llm_json(llm, prompt_text, retries=2):
    for i in range(retries + 1):
        resp = llm.invoke(prompt_text)
        text = resp.content.strip()
        json_part = re.search(r'\{[\s\S]*\}', text)
        if not json_part:
            continue
        try:
            return json.loads(json_part.group(0))
        except json.JSONDecodeError:
            prompt_text += "\n\nJSON 형식만 정확히 출력해주세요."
    raise ValueError("❌ LLM이 올바른 JSON을 반환하지 못했습니다.")

# ===================== 회의 일자 추정 프롬프트 =====================
meeting_date_prompt = PromptTemplate.from_template("""
다음 회의 대화 내용을 보고 회의가 실제로 열린 날짜를 추정하세요.
상대적 표현(오늘, 내일, 이번 주, 다음 주 등)을 고려하여 ISO 형식으로 작성합니다.

출력 예시:
{{ "meeting_date": "2025-10-27" }}

회의록:
{text}
""")

# ===================== 핵심 파이프라인 =====================
def run_meeting_pipeline(audio_path: str) -> dict:
    if not os.path.exists(audio_path):
        raise FileNotFoundError(f"❌ 파일 없음: {audio_path}")

    # === 1️⃣ Whisper STT 변환 ===
    model = whisper.load_model("small")
    print(f"🎙️ Whisper 변환 중... {audio_path}")
    result = model.transcribe(audio_path, language="ko")
    full_text = result["text"].strip()

    # === 2️⃣ 회의일자 추정 ===
    llm_date = ChatOpenAI(model_name="gpt-4o-mini", temperature=0)
    try:
        date_json = safe_llm_json(llm_date, meeting_date_prompt.format(text=full_text))
        base_dt = dateparser.parse(date_json.get("meeting_date", ""), languages=["ko"]) or datetime.now()

        # 🔧 연도 보정 (LLM이 과거 연도 추정할 경우)
        if base_dt.year < datetime.now().year:
            base_dt = base_dt.replace(year=datetime.now().year)

    except Exception:
        base_dt = datetime.now()


    # === 3️⃣ 회의 요약 / 결정사항 / 액션아이템 추출 ===
    llm = ChatOpenAI(model_name="gpt-4o-mini", temperature=0.2)
    prompt_text = meeting_summary_prompt.format(text=full_text)
    parsed_json = safe_llm_json(llm, prompt_text)

    # === 4️⃣ due 날짜 정규화 (문맥 기반 변환) ===
    for item in parsed_json.get("action_items", []):
        item["due"] = normalize_due(item.get("due"), base_dt)

    # === 5️⃣ fallback: due가 전부 None이면 순차 배정 ===
    for idx, item in enumerate(parsed_json.get("action_items", [])):
        if not item.get("due"):
            item["due"] = (base_dt + timedelta(days=idx)).strftime("%Y-%m-%d")

    # === 6️⃣ Pydantic 검증 ===
    validated = MeetingSummary(**parsed_json)

        # === 7️⃣ DB 저장 (내 DB + 팀원 DB) ===
    try:
        # ✅ 기존 개인 DB (외부 서버)
        conn1 = pymysql.connect(
            host=DB_HOST, user=DB_USER, password=DB_PASSWORD,
            database=DB_NAME, port=DB_PORT, charset="utf8mb4"
        )
        cur1 = conn1.cursor()
        cur1.execute("""
        INSERT INTO meeting_summary (meeting_file, topic_summary, content_summary, decisions, action_items)
        VALUES (%s,%s,%s,%s,%s)
        ON DUPLICATE KEY UPDATE
            topic_summary=VALUES(topic_summary),
            content_summary=VALUES(content_summary),
            decisions=VALUES(decisions),
            action_items=VALUES(action_items),
            created_at=CURRENT_TIMESTAMP;
        """, (
            audio_path,
            validated.topic_summary,
            validated.content_summary,
            json.dumps(validated.decisions, ensure_ascii=False),
            json.dumps([a.dict() for a in validated.action_items], ensure_ascii=False)
        ))
        conn1.commit()
        print("✅ 개인 DB 저장 완료")

    except Exception as e:
        print("❌ 개인 DB 오류:", e)
    finally:
        conn1.close()

    # === 🧩 팀원 DB에도 추가 저장 ===
    DB_CONFIG = {
        'host': 'localhost',
        'user': 'admin',
        'password': '1qazZAQ!',
        'db': 'final',
        'charset': 'utf8mb4'
    }

    try:
        conn2 = pymysql.connect(**DB_CONFIG)
        cur2 = conn2.cursor()

        # ⚙️ 테이블 자동 생성 (없을 시)
        cur2.execute("""
        CREATE TABLE IF NOT EXISTS team_meeting_summary (
            id INT AUTO_INCREMENT PRIMARY KEY,
            meeting_file VARCHAR(255) UNIQUE,
            topic_summary TEXT,
            content_summary TEXT,
            decisions JSON,
            action_items JSON,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        ) CHARACTER SET utf8mb4;
        """)

        # 💾 데이터 삽입 (중복 방지)
        cur2.execute("""
        INSERT INTO team_meeting_summary (meeting_file, topic_summary, content_summary, decisions, action_items)
        VALUES (%s,%s,%s,%s,%s)
        ON DUPLICATE KEY UPDATE
            topic_summary=VALUES(topic_summary),
            content_summary=VALUES(content_summary),
            decisions=VALUES(decisions),
            action_items=VALUES(action_items),
            created_at=CURRENT_TIMESTAMP;
        """, (
            audio_path,
            validated.topic_summary,
            validated.content_summary,
            json.dumps(validated.decisions, ensure_ascii=False),
            json.dumps([a.dict() for a in validated.action_items], ensure_ascii=False)
        ))
        conn2.commit()
        print("✅ 팀원 DB 저장 완료")

    except Exception as e:
        print("❌ 팀원 DB 오류:", e)
    finally:
        conn2.close()

    # === 9️⃣ DOCX 파일 자동 생성 ===
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt
    from docx.oxml.ns import qn   # ✅ 이 줄 추가


    doc_dir = "static/docs"
    os.makedirs(doc_dir, exist_ok=True)
    base_filename = os.path.splitext(os.path.basename(audio_path))[0]  # 확장자 제거
    doc_path = os.path.join(
        doc_dir,
        f"회의록_{base_dt.strftime('%Y-%m-%d')}_{base_filename}.docx"
    )

    doc = Document()
    # 스타일 지정
    style = doc.styles['Normal']
    style.font.name = 'Malgun Gothic'  # 🔹 윈도우에서 존재하는 한글 폰트
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
    style.font.size = Pt(12)

    # --- 제목 ---
    title = doc.add_heading("회의 요약 보고서", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # --- 기본 정보 ---
    doc.add_paragraph(f"📅 회의일자: {base_dt.strftime('%Y-%m-%d')}")
    doc.add_paragraph(f"🎧 파일명: {base_filename}")
    doc.add_paragraph("")

    # --- 주제 요약 ---
    doc.add_heading("1. 주제 요약", level=2)
    doc.add_paragraph(validated.topic_summary)

    # --- 내용 요약 ---
    doc.add_heading("2. 내용 요약", level=2)
    doc.add_paragraph(validated.content_summary)

    # --- 결정사항 ---
    doc.add_heading("3. 결정사항", level=2)
    for d in validated.decisions:
        doc.add_paragraph(f"• {d}", style="List Bullet")

    # --- 액션 아이템 ---
    doc.add_heading("4. 액션 아이템", level=2)
    for item in validated.action_items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(f"담당자: {item.name}\n").bold = True
        p.add_run(f"작업내용: {item.task}\n")
        p.add_run(f"기한: {item.due if item.due else '미정'}")
        
        # === 8️⃣ JSON 파일도 자동 저장 (프론트에서 보기용) ===
    json_dir = "static/data"
    os.makedirs(json_dir, exist_ok=True)
    json_path = os.path.join(json_dir, f"{base_filename}.json")

    with open(json_path, "w", encoding="utf-8") as f:
        json.dump({
            "topic_summary": validated.topic_summary,
            "content_summary": validated.content_summary,
            "decisions": validated.decisions,
            "action_items": [a.dict() for a in validated.action_items],
        }, f, ensure_ascii=False, indent=2)

    print(f"📄 JSON 저장 완료: {json_path}")

    doc.save(doc_path)
    print(f"📝 DOCX 저장 완료: {doc_path}")
     # === 10️⃣ 결과 반환 ===
    return {
        "topic_summary": validated.topic_summary,
        "content_summary": validated.content_summary,
        "decisions": validated.decisions,
        "action_items": [a.dict() for a in validated.action_items],
        "docx_path": doc_path
    }

